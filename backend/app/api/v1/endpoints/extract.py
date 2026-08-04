"""
Document text extraction endpoint.
Supports: .docx, .txt, .md, .pdf, .png, .jpg, .jpeg, .webp
Images are processed via OCR.space free API with auto-resize for large files.
"""
import logging
import io
from fastapi import APIRouter, UploadFile, File, HTTPException
from pydantic import BaseModel

logger = logging.getLogger(__name__)
router = APIRouter()

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
IMAGE_MIME_TYPES = {"image/png", "image/jpeg", "image/jpg", "image/webp", "image/gif"}


class ExtractResponse(BaseModel):
    filename: str
    text: str
    char_count: int


def _is_image(filename: str, content_type: str) -> bool:
    import os
    ext = os.path.splitext(filename.lower())[1]
    return ext in IMAGE_EXTENSIONS or content_type in IMAGE_MIME_TYPES


@router.post("/", response_model=ExtractResponse, summary="Extract text from an uploaded document or image")
async def extract_document(file: UploadFile = File(...)) -> ExtractResponse:
    filename = file.filename or "unknown"
    content_type = file.content_type or ""
    content = await file.read()

    text = ""

    try:
        # ── Images — use OCR.space API ─────────────────────────────────────
        if _is_image(filename, content_type):
            text = await _extract_image_text(content, content_type or "image/jpeg", filename)

        # ── .docx ──────────────────────────────────────────────────────────
        elif filename.endswith(".docx") or content_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            try:
                import docx
                doc = docx.Document(io.BytesIO(content))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text.strip())
                text = "\n".join(paragraphs)
            except ImportError:
                raise HTTPException(status_code=500, detail="python-docx is not installed on the server.")

        # ── .pdf ───────────────────────────────────────────────────────────
        elif filename.endswith(".pdf") or content_type == "application/pdf":
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    pages = [page.extract_text() for page in pdf.pages if page.extract_text()]
                text = "\n\n".join(pages)
            except ImportError:
                raise HTTPException(status_code=500, detail="PDF extraction library not available.")

        # ── .txt / .md / .json ─────────────────────────────────────────────
        elif (
            filename.endswith(".txt")
            or filename.endswith(".md")
            or filename.endswith(".json")
            or content_type.startswith("text/")
        ):
            text = content.decode("utf-8", errors="replace")

        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {filename}. Supported: .docx, .pdf, .txt, .md, .png, .jpg, .jpeg"
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Failed to extract from %s", filename)
        raise HTTPException(status_code=500, detail=f"Extraction failed: {str(e)}")

    text = text.strip()[:8000]

    if not text:
        raise HTTPException(
            status_code=422,
            detail="No text could be extracted from the file."
        )

    return ExtractResponse(filename=filename, text=text, char_count=len(text))


async def _extract_image_text(content: bytes, content_type: str, filename: str) -> str:
    """
    Extract text from an image using OCR.space free API.
    Auto-resizes images over 900KB using Pillow to stay under the 1MB free tier limit.
    """
    import httpx
    import os
    from app.core.config import settings

    ocr_api_key = getattr(settings, "ocr_api_key", "") or "helloworld"
    ext = os.path.splitext(filename.lower())[1].lstrip(".") or "png"
    filetype_map = {"jpg": "JPG", "jpeg": "JPG", "png": "PNG", "gif": "GIF", "webp": "JPG"}
    filetype = filetype_map.get(ext, "PNG")
    ocr_content = content
    ocr_filename = filename
    ocr_mime = content_type or f"image/{ext}"

    # ── Resize if over 900KB ───────────────────────────────────────────────
    if len(content) > 900_000:
        try:
            from PIL import Image as PILImage
            img = PILImage.open(io.BytesIO(content))
            if img.mode in ("RGBA", "P", "LA"):
                img = img.convert("RGB")
            quality = 85
            scale = 0.9
            while True:
                w, h = img.size
                resized = img.resize((int(w * scale), int(h * scale)), PILImage.LANCZOS)
                buf = io.BytesIO()
                resized.save(buf, format="JPEG", quality=quality)
                ocr_content = buf.getvalue()
                if len(ocr_content) <= 900_000 or scale < 0.2:
                    break
                scale *= 0.85
                quality = max(60, quality - 5)
            filetype = "JPG"
            ocr_filename = os.path.splitext(filename)[0] + ".jpg"
            ocr_mime = "image/jpeg"
            logger.info("Resized image from %d to %d bytes for OCR", len(content), len(ocr_content))
        except ImportError:
            logger.warning("Pillow not available — sending original image, may exceed limit")

    # ── Try Engine 1 first (works for all sizes) ───────────────────────────
    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.post(
            "https://api.ocr.space/parse/image",
            data={
                "apikey": ocr_api_key,
                "language": "eng",
                "isOverlayRequired": "false",
                "detectOrientation": "true",
                "scale": "true",
                "OCREngine": "1",
                "filetype": filetype,
            },
            files={"file": (ocr_filename, ocr_content, ocr_mime)},
        )

    if response.status_code == 413:
        raise HTTPException(
            status_code=413,
            detail="Image is too large for OCR processing (max 1MB). Please use a smaller image."
        )
    if response.status_code != 200:
        raise HTTPException(status_code=500, detail=f"OCR API returned {response.status_code}")

    result = response.json()

    # ── Fallback to Engine 2 if Engine 1 fails ─────────────────────────────
    if result.get("IsErroredOnProcessing"):
        err = " ".join(str(m) for m in result.get("ErrorMessage", []))
        logger.warning("OCR Engine 1 failed for %s: %s — trying Engine 2", filename, err)
        async with httpx.AsyncClient(timeout=60) as client:
            response = await client.post(
                "https://api.ocr.space/parse/image",
                data={
                    "apikey": ocr_api_key,
                    "language": "eng",
                    "isOverlayRequired": "false",
                    "detectOrientation": "true",
                    "scale": "true",
                    "OCREngine": "2",
                    "filetype": filetype,
                },
                files={"file": (ocr_filename, ocr_content, ocr_mime)},
            )
        result = response.json()

    if result.get("IsErroredOnProcessing"):
        err = " ".join(str(m) for m in result.get("ErrorMessage", []))
        logger.error("OCR failed for %s: %s", filename, err)
        raise HTTPException(status_code=500, detail=f"OCR failed: {err}")

    parsed_results = result.get("ParsedResults", [])
    if not parsed_results:
        return "[No text found in image]"

    texts = [r.get("ParsedText", "") for r in parsed_results if r.get("ParsedText")]
    extracted = "\n".join(texts).strip()

    if not extracted:
        return "[No readable text found in image. The image may not contain printed text.]"

    logger.info("OCR extraction successful for %s (%d chars)", filename, len(extracted))
    return extracted
